# -*- coding: utf-8 -*-
"""
Script tạo file Word khảo sát AURA ARCHIVE
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Tạo document
doc = Document()

# ===== STYLES =====
# Title style
title_style = doc.styles['Title']
title_style.font.size = Pt(24)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)

# Heading 1
h1_style = doc.styles['Heading 1']
h1_style.font.size = Pt(16)
h1_style.font.bold = True

# Heading 2
h2_style = doc.styles['Heading 2']
h2_style.font.size = Pt(14)
h2_style.font.bold = True

# ===== CONTENT =====

# Title
title = doc.add_heading('BỘ CÂU HỎI KHẢO SÁT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('DỰ ÁN AURA ARCHIVE')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph('Nền tảng E-commerce Thời trang Cao cấp').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Mục đích
purpose = doc.add_paragraph()
purpose.add_run('Mục đích: ').bold = True
purpose.add_run('Thu thập ý kiến người dùng tiềm năng để hoàn thiện nền tảng e-commerce thời trang cao cấp AURA ARCHIVE')

doc.add_paragraph()
doc.add_paragraph('─' * 60)
doc.add_paragraph()

# ===== PHẦN I =====
doc.add_heading('PHẦN I: THÔNG TIN CƠ BẢN', level=1)

doc.add_heading('1. Thói quen mua sắm', level=2)

doc.add_paragraph('Q1.1. Bạn thường mua sắm thời trang online với tần suất như thế nào?')
options = ['☐ Hàng tuần', '☐ 2-3 lần/tháng', '☐ 1 lần/tháng', '☐ Vài lần/năm', '☐ Hiếm khi']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q1.2. Bạn đã từng mua sản phẩm thời trang luxury/designer chưa?')
options = ['☐ Có, thường xuyên', '☐ Có, thỉnh thoảng', '☐ Có, nhưng rất ít', '☐ Chưa bao giờ, nhưng quan tâm', '☐ Chưa và không quan tâm']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q1.3. Ngân sách trung bình cho 1 món đồ thời trang của bạn?')
options = ['☐ Dưới 500.000 VNĐ', '☐ 500.000 - 2.000.000 VNĐ', '☐ 2.000.000 - 5.000.000 VNĐ', '☐ 5.000.000 - 10.000.000 VNĐ', '☐ Trên 10.000.000 VNĐ']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN II =====
doc.add_heading('PHẦN II: ĐÁNH GIÁ CÁC TÍNH NĂNG CHÍNH', level=1)

doc.add_paragraph('Hướng dẫn: Đánh giá mức độ quan trọng của mỗi tính năng (1-5 sao)')
doc.add_paragraph('1 = Không quan trọng | 5 = Rất quan trọng')

doc.add_heading('2. Tính năng mua hàng cơ bản', level=2)

# Tạo bảng đánh giá
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'

# Header
headers = ['Tính năng', '1 ⭐', '2 ⭐', '3 ⭐', '4 ⭐', '5 ⭐']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

# Features
features = [
    'Q2.1. Bộ lọc sản phẩm (thương hiệu, loại, giá, màu)',
    'Q2.2. Xem sản phẩm đã xem gần đây',
    'Q2.3. So sánh sản phẩm',
    'Q2.4. Hướng dẫn chọn size',
    'Q2.5. Wishlist (Danh sách yêu thích)',
    'Q2.6. Đánh giá & Review sản phẩm'
]
for i, feature in enumerate(features):
    table.rows[i+1].cells[0].text = feature

doc.add_paragraph()
doc.add_heading('3. Thanh toán & Vận chuyển', level=2)

doc.add_paragraph('Q3.1. Phương thức thanh toán nào bạn ưu tiên? (Chọn tối đa 3)')
options = ['☐ MoMo', '☐ VNPay', '☐ Chuyển khoản ngân hàng (QR Code)', '☐ COD (Thanh toán khi nhận hàng)', '☐ Thẻ tín dụng/ghi nợ quốc tế', '☐ PayPal']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q3.2. Mức độ quan trọng của việc tính phí ship tự động theo địa chỉ?')
options = ['☐ 1 - Không quan trọng', '☐ 2 - Ít quan trọng', '☐ 3 - Trung bình', '☐ 4 - Quan trọng', '☐ 5 - Rất quan trọng']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q3.3. Bạn có sử dụng mã giảm giá/coupon khi mua hàng online không?')
options = ['☐ Luôn luôn tìm trước khi mua', '☐ Có, nếu có sẵn', '☐ Thỉnh thoảng', '☐ Hiếm khi', '☐ Không bao giờ']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN III =====
doc.add_heading('PHẦN III: TƯ VẤN AI (AI STYLIST)', level=1)

note = doc.add_paragraph()
note.add_run('Mô tả tính năng: ').bold = True
note.add_run('Chatbot AI tư vấn thời trang, giúp gợi ý sản phẩm phù hợp với phong cách và dịp sử dụng')

doc.add_paragraph()
doc.add_paragraph('Q4.1. Bạn có hứng thú với tính năng AI Stylist tư vấn thời trang không?')
options = ['☐ Rất hứng thú, muốn sử dụng ngay', '☐ Hứng thú, sẽ thử nghiệm', '☐ Bình thường, có thể sử dụng', '☐ Không quan tâm lắm', '☐ Không cần thiết']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q4.2. Bạn mong muốn AI Stylist có khả năng nào nhất? (Chọn tối đa 3)')
options = ['☐ Gợi ý outfit phù hợp với dịp (đi làm, dạo phố, party...)', '☐ Phân tích phong cách cá nhân', '☐ Gợi ý sản phẩm theo ngân sách', '☐ Phối đồ với những item đã có', '☐ Tư vấn về thương hiệu và chất liệu', '☐ Theo dõi xu hướng thời trang mới']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q4.3. Ngôn ngữ nào bạn muốn AI Stylist hỗ trợ?')
options = ['☐ Chỉ tiếng Việt là đủ', '☐ Tiếng Việt và tiếng Anh', '☐ Cần thêm các ngôn ngữ khác: ____________']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN IV =====
doc.add_heading('PHẦN IV: TRẢI NGHIỆM GIAO DIỆN (UI/UX)', level=1)

doc.add_paragraph('Q5.1. Bạn thường sử dụng thiết bị nào để mua sắm online?')
options = ['☐ Điện thoại (70%+)', '☐ Laptop/Desktop (70%+)', '☐ Cả hai, tùy lúc']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q5.2. Yếu tố nào quan trọng nhất với bạn khi mua sắm online? (Xếp hạng 1-5)')
options = ['___ Tốc độ tải trang nhanh', '___ Giao diện đẹp, sang trọng', '___ Dễ tìm kiếm sản phẩm', '___ Thông tin sản phẩm chi tiết', '___ Quy trình thanh toán đơn giản']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q5.3. Bạn thích giao diện website thời trang kiểu nào?')
options = ['☐ Tối giản (Minimalist - ít màu, nhiều khoảng trắng)', '☐ Sang trọng (Luxury - tone màu đen/vàng)', '☐ Trẻ trung (Colorful - nhiều màu sắc)', '☐ Hiện đại (Modern - gradient, animation mượt)']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q5.4. Bạn có cần website hỗ trợ đa ngôn ngữ (Tiếng Việt/English) không?')
options = ['☐ Có, rất cần thiết', '☐ Có thể có càng tốt', '☐ Không cần thiết']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN V =====
doc.add_heading('PHẦN V: QUẢN LÝ TÀI KHOẢN', level=1)

doc.add_paragraph('Q6.1. Bạn ưu tiên đăng nhập bằng phương thức nào?')
options = ['☐ Email/Password truyền thống', '☐ Google account', '☐ Facebook account', '☐ OTP qua số điện thoại', '☐ Tất cả đều ok']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q6.2. Tính năng nào bạn muốn có trong trang tài khoản cá nhân? (Chọn nhiều)')
options = ['☐ Lịch sử đơn hàng chi tiết', '☐ Theo dõi trạng thái đơn hàng real-time', '☐ Quản lý nhiều địa chỉ giao hàng', '☐ Danh sách yêu thích (Wishlist)', '☐ Lịch sử sản phẩm đã xem', '☐ Điểm thưởng/Loyalty points', '☐ Thông báo khuyến mãi']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN VI =====
doc.add_heading('PHẦN VI: TÍNH NĂNG NÂNG CAO', level=1)

doc.add_paragraph('Q7.1. Bạn có muốn nhận thông báo về sản phẩm mới/giảm giá qua các kênh nào? (Chọn nhiều)')
options = ['☐ Email newsletter', '☐ SMS', '☐ Push notification (app/browser)', '☐ Không muốn nhận']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q7.2. Tính năng nào sau đây bạn thấy hữu ích cho website thời trang? (Chọn tối đa 3)')
options = ['☐ Blog thời trang & tips phối đồ', '☐ Virtual try-on (Thử đồ ảo)', '☐ Tìm kiếm bằng hình ảnh', '☐ Ký gửi/Consignment (Bán lại đồ cũ)', '☐ Đặt lịch hẹn tư vấn 1-1', '☐ Loyalty program với ưu đãi đặc biệt']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q7.3. Bạn có quan tâm đến tính năng KÝ GỬI (Consignment) - bán lại đồ designer đã sử dụng không?')
options = ['☐ Rất quan tâm, muốn sử dụng', '☐ Quan tâm, cần biết thêm thông tin', '☐ Có thể sử dụng trong tương lai', '☐ Không quan tâm']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN VII =====
doc.add_heading('PHẦN VII: BẢO MẬT & TIN CẬY', level=1)

doc.add_paragraph('Q8.1. Yếu tố nào khiến bạn tin tưởng một website thời trang? (Chọn tối đa 3)')
options = ['☐ Chứng nhận hàng chính hãng/Authenticated', '☐ Chính sách đổi trả rõ ràng', '☐ Reviews từ người mua thực', '☐ Bảo mật thanh toán (SSL, cổng thanh toán uy tín)', '☐ Thương hiệu/tên tuổi của shop', '☐ Responses nhanh từ customer service']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q8.2. Bạn có lo ngại về vấn đề bảo mật thông tin cá nhân khi mua hàng online không?')
options = ['☐ Rất lo ngại', '☐ Lo ngại một chút', '☐ Bình thường', '☐ Không lo ngại']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN VIII =====
doc.add_heading('PHẦN VIII: PHẢN HỒI MỞ', level=1)

doc.add_paragraph('Q9.1. Bạn gặp vấn đề/khó chịu gì nhất khi mua sắm thời trang online?')
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)

doc.add_paragraph()
doc.add_paragraph('Q9.2. Có tính năng nào bạn mong muốn mà chưa được đề cập ở trên không?')
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)

doc.add_paragraph()
doc.add_paragraph('Q9.3. Bạn có góp ý nào khác cho dự án AURA ARCHIVE không?')
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ===== PHẦN IX =====
doc.add_heading('PHẦN IX: THÔNG TIN KHÁCH HÀNG', level=1)

note = doc.add_paragraph()
note.add_run('Lưu ý: ').bold = True
note.add_run('Thông tin này giúp chúng tôi phân tích và cải thiện dịch vụ. Các trường có dấu (*) là bắt buộc.')

doc.add_heading('Thông tin cơ bản', level=2)

doc.add_paragraph('Q10.1. Độ tuổi của bạn? (*)')
options = ['☐ Dưới 18 tuổi', '☐ 18 - 24 tuổi', '☐ 25 - 34 tuổi', '☐ 35 - 44 tuổi', '☐ 45 - 54 tuổi', '☐ Trên 54 tuổi']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q10.2. Giới tính? (*)')
options = ['☐ Nam', '☐ Nữ', '☐ Khác / Không muốn trả lời']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q10.3. Tình trạng hôn nhân?')
options = ['☐ Độc thân', '☐ Đã kết hôn / Có gia đình', '☐ Khác']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_heading('Nghề nghiệp & Thu nhập', level=2)

doc.add_paragraph('Q10.4. Nghề nghiệp hiện tại của bạn? (*)')
options = ['☐ Học sinh / Sinh viên', '☐ Nhân viên văn phòng', '☐ Quản lý / Giám đốc', '☐ Tự kinh doanh / Chủ doanh nghiệp', '☐ Freelancer / Làm việc tự do', '☐ Nghệ sĩ / Người mẫu / Influencer', '☐ Công chức / Viên chức nhà nước', '☐ Nội trợ', '☐ Đã nghỉ hưu', '☐ Khác: _______________']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q10.5. Mức thu nhập hàng tháng? (Không bắt buộc)')
options = ['☐ Dưới 10 triệu VNĐ', '☐ 10 - 20 triệu VNĐ', '☐ 20 - 40 triệu VNĐ', '☐ 40 - 70 triệu VNĐ', '☐ Trên 70 triệu VNĐ', '☐ Không muốn trả lời']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_heading('Khu vực & Thông tin liên hệ', level=2)

doc.add_paragraph('Q10.6. Bạn đang sinh sống tại khu vực nào? (*)')
options = ['☐ TP. Hồ Chí Minh', '☐ Hà Nội', '☐ Đà Nẵng', '☐ Các tỉnh miền Bắc khác', '☐ Các tỉnh miền Trung khác', '☐ Các tỉnh miền Nam khác', '☐ Nước ngoài']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q10.7. Bạn biết đến khảo sát này qua kênh nào?')
options = ['☐ Facebook', '☐ Instagram', '☐ TikTok', '☐ Zalo', '☐ Email', '☐ Bạn bè giới thiệu', '☐ Website/Blog', '☐ Khác: _______________']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_heading('Thông tin liên hệ (Không bắt buộc)', level=2)

# Bảng thông tin liên hệ
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Thông tin'
table.rows[0].cells[1].text = 'Điền vào'
table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
table.rows[1].cells[0].text = 'Họ và tên'
table.rows[2].cells[0].text = 'Email'
table.rows[3].cells[0].text = 'Số điện thoại'

doc.add_paragraph()
doc.add_paragraph('Q10.8. Bạn có muốn nhận thông tin về AURA ARCHIVE sau khi ra mắt không?')
options = ['☐ Có, qua Email', '☐ Có, qua SMS/Zalo', '☐ Có, cả Email và SMS', '☐ Không, cảm ơn']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Q10.9. Bạn có sẵn sàng tham gia phỏng vấn sâu (30 phút, có quà tặng) không?')
options = ['☐ Có, sẵn sàng', '☐ Có thể, tùy thời gian', '☐ Không']
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('─' * 60)
doc.add_paragraph()

# Lời cảm ơn
thanks = doc.add_paragraph()
thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
thanks.add_run('🙏 CẢM ƠN BẠN ĐÃ DÀNH THỜI GIAN HOÀN THÀNH KHẢO SÁT! 🙏').bold = True

end_note = doc.add_paragraph('Phản hồi của bạn sẽ giúp chúng tôi xây dựng một nền tảng thời trang tốt hơn.')
end_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'C:\Users\anhvu\OneDrive\Desktop\KLTN\AURA_ARCHIVE_Survey.docx'
doc.save(output_path)
print(f'✅ Đã tạo file Word thành công: {output_path}')
